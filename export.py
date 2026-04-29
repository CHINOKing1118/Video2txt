import os
import re

from utils import retry_operation, write_text_file


def md_to_docx(md_text, docx_path):
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise Exception("python-docx 未安装，请运行: pip install python-docx")

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)

    def add_formatted_text(paragraph, text):
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    for line in md_text.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, stripped[2:])
        elif re.match(r'^\d+\.\s', stripped):
            content = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, content)
        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            add_formatted_text(p, stripped[2:])
            p.paragraph_format.left_indent = Pt(20)
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, stripped)
    doc.save(docx_path)


def md_to_pdf(md_text, pdf_path):
    try:
        from fpdf import FPDF
    except ImportError:
        raise Exception("fpdf2 未安装，请运行: pip install fpdf2")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    font_name = None
    user_font_dir = os.path.join(os.environ.get("LOCALAPPDATA", ""),
                                 "Microsoft", "Windows", "Fonts")
    font_candidates = [
        (os.path.join(user_font_dir, "MiSans-Regular.ttf"), "MiSans"),
        (os.path.join(user_font_dir, "MiSans-Normal.ttf"), "MiSans"),
        ("C:/Windows/Fonts/simhei.ttf", "SimHei"),
        ("C:/Windows/Fonts/msyh.ttc", "MsYh"),
        ("C:/Windows/Fonts/simsun.ttc", "SimSun"),
    ]
    for fp, name in font_candidates:
        if os.path.exists(fp):
            try:
                pdf.add_font(name, "", fp)
                font_name = name
                break
            except Exception:
                continue

    def set_size(size):
        if font_name:
            pdf.set_font(font_name, size=size)
        else:
            pdf.set_font("Helvetica", size=size)

    def safe_write(line_text, line_height=6):
        try:
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, line_height, line_text)
        except Exception:
            for ch in line_text:
                try:
                    pdf.set_x(pdf.l_margin)
                    pdf.multi_cell(0, line_height, ch)
                except Exception:
                    pass

    set_size(11)
    for line in md_text.split('\n'):
        stripped = line.strip()
        if not stripped:
            pdf.ln(3)
            continue
        clean = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
        if clean.startswith('### '):
            set_size(13)
            safe_write(clean[4:], 7)
            set_size(11)
        elif clean.startswith('## '):
            set_size(15)
            safe_write(clean[3:], 8)
            set_size(11)
        elif clean.startswith('# '):
            set_size(18)
            safe_write(clean[2:], 10)
            set_size(11)
        elif clean.startswith('- ') or clean.startswith('* '):
            safe_write(f"  \u2022 {clean[2:]}")
        elif re.match(r'^\d+\.\s', clean):
            safe_write(clean)
        elif clean.startswith('> '):
            safe_write(f"  {clean[2:]}")
        else:
            safe_write(clean)
    pdf.output(pdf_path)


def format_subtitle(segments):
    lines = []
    for seg in segments:
        start = seg.get("start", 0)
        text = seg.get("text", "").strip()
        if not text:
            continue
        m, s = divmod(int(start), 60)
        h, m = divmod(m, 60)
        lines.append(f"[{h:02d}:{m:02d}:{s:02d}] {text}")
    return "\n".join(lines)


def save_output(pol_text, folder_path, fmt):
    if fmt == "docx":
        path = os.path.join(folder_path, "文稿.docx")
        retry_operation(lambda: md_to_docx(pol_text, path))
    elif fmt == "pdf":
        path = os.path.join(folder_path, "文稿.pdf")
        retry_operation(lambda: md_to_pdf(pol_text, path))
    else:
        path = os.path.join(folder_path, "文稿.md")
        retry_operation(lambda: write_text_file(path, pol_text))
    return path
